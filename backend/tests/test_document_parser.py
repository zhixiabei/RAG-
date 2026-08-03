import unittest
from io import BytesIO
import struct
from unittest.mock import patch
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document
from rag_app.infrastructure.parsing.document_parser import CHUNK_OVERLAP, CHUNK_SIZE, DocumentParser


class DocumentParserSpecialExtensionTest(unittest.TestCase):
    def setUp(self):
        self.parser = DocumentParser()

    def test_pptx_stream_reads_slide_text_in_order_without_loading_media(self):
        presentation = b"""<?xml version="1.0" encoding="UTF-8"?>
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst><p:sldId id="256" r:id="rId2"/><p:sldId id="257" r:id="rId1"/></p:sldIdLst>
</p:presentation>"""
        relationships = b"""<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide2.xml"/>
</Relationships>"""

        def slide_xml(text):
            return f"""<?xml version="1.0" encoding="UTF-8"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>{text}</a:t></a:r></a:p></p:txBody></p:sp></p:spTree></p:cSld>
</p:sld>""".encode()

        stream = BytesIO()
        with ZipFile(stream, "w") as archive:
            archive.writestr("ppt/presentation.xml", presentation)
            archive.writestr("ppt/_rels/presentation.xml.rels", relationships)
            archive.writestr("ppt/slides/slide1.xml", slide_xml("second"))
            archive.writestr("ppt/slides/slide2.xml", slide_xml("first"))
            archive.writestr("ppt/media/video1.mp4", b"x" * (2 * 1024 * 1024))
        stream.seek(0)

        chunks = self.parser.parse_stream("large-media.pptx", stream)

        self.assertEqual(len(chunks), 1)
        self.assertIsNone(chunks[0].page_number)
        self.assertIn("first\n\nsecond", chunks[0].text)

    def test_att_text_file_remains_parsable_but_is_not_uploadable(self):
        chunks = self.parser.parse("长63渗透率.att", "渗透率 12.5 mD".encode("utf-8"))

        self.assertFalse(self.parser.supports("长63渗透率.att"))
        self.assertTrue(all(len(chunk.text) <= CHUNK_SIZE for chunk in chunks))
        self.assertIn("渗透率 12.5 mD", chunks[0].text)

    def test_docx_ignores_relationship_to_missing_optional_custom_ui_part(self):
        document_stream = BytesIO()
        document = Document()
        document.add_paragraph("化163-1井示踪剂施工设计")
        document.save(document_stream)

        damaged_stream = BytesIO()
        with ZipFile(BytesIO(document_stream.getvalue())) as source, ZipFile(damaged_stream, "w") as damaged:
            for entry in source.infolist():
                data = source.read(entry)
                if entry.filename == "_rels/.rels":
                    root = ElementTree.fromstring(data)
                    namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
                    ElementTree.SubElement(
                        root,
                        f"{{{namespace}}}Relationship",
                        {
                            "Id": "rIdMissingCustomUi",
                            "Type": "http://schemas.microsoft.com/office/2006/relationships/ui/extensibility",
                            "Target": "userCustomization/customUI.xml",
                        },
                    )
                    data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
                damaged.writestr(entry, data)

        damaged_content = damaged_stream.getvalue()
        with self.assertRaisesRegex(KeyError, "userCustomization/customUI.xml"):
            Document(BytesIO(damaged_content))

        chunks = self.parser.parse("化163-1井示踪剂施工设计.docx", damaged_content)

        self.assertIn("化163-1井示踪剂施工设计", "\n".join(chunk.text for chunk in chunks))

    def test_geomap_v360_layer_style_is_parsed_as_structured_records(self):
        header = bytearray(516)
        header[:24] = b"Geomap v3.60 LayerStyle\x00"
        struct.pack_into("<I", header, 512, 1)
        record = bytearray(524)
        values = ["渗透率填色", "符号", "符号大小", "前景色", "背景色", "边界光滑", "光滑系数"]
        slots = ((0, 64), (80, 64), (160, 64), (240, 64), (320, 64), (400, 64), (480, 32))
        for value, (offset, width) in zip(values, slots):
            encoded = value.encode("gb18030")
            self.assertLess(len(encoded), width)
            record[offset:offset + len(encoded)] = encoded

        chunks = self.parser.parse("长63渗透率.att", bytes(header + record))
        text = "\n".join(chunk.text for chunk in chunks)

        self.assertIn("Geomap v3.60 LayerStyle", text)
        self.assertIn("LayerStyle 记录数: 1", text)
        self.assertIn("属性数据表", text)
        self.assertEqual(len(chunks), 1)
        self.assertIn("渗透率填色", text)
        self.assertIn("符号大小", text)
        self.assertIn("光滑系数", text)

    def test_geomap_v360_map_extracts_layer_attribute_tables_with_standard_chunking(self):
        content = bytearray(2048)
        content[:17] = b"Geomap v3.60 Map\x00"
        map_name = "长611渗透率".encode("gb18030")
        content[0x204:0x204 + len(map_name)] = map_name

        def append_layer(name, values):
            header = bytearray(2048)
            header[:18] = b"Geomap v3.60 Layer"
            encoded_name = name.encode("gb18030")
            struct.pack_into("<I", header, 0x362, len(encoded_name))
            header[0x366:0x366 + len(encoded_name)] = encoded_name
            content.extend(header)
            for value in values:
                encoded_value = value.encode("gb18030")
                content.extend(b"\x2f\x3d\x50\x75")
                content.extend(struct.pack("<I", len(encoded_value)))
                content.extend(encoded_value)
                content.extend(b"\x1e\x00\x00\x00")

        well_names = [f"化{i}-1" for i in range(1, 301)]
        append_layer("井位", well_names)
        append_layer("渗透率", ["0.2", "0.4", "1.3"])

        chunks = self.parser.parse("长611渗透率.gdb", bytes(content))
        text = "\n".join(chunk.text for chunk in chunks)

        self.assertGreater(len(chunks), 3)
        self.assertTrue(all(len(chunk.text) <= CHUNK_SIZE for chunk in chunks))
        self.assertIn("格式: Geomap v3.60 Map 属性数据表", text)
        self.assertIn("图层名称: 井位", text)
        self.assertIn("300 化300-1", text)
        self.assertIn("图层名称: 渗透率", text)
        self.assertIn("0.2", text)
        self.assertNotIn("二进制内容", text)

    def test_geomap_v360_map_accepts_final_metadata_only_short_layer(self):
        content = bytearray(2048)
        content[:17] = b"Geomap v3.60 Map\x00"
        short_layer = bytearray(916)
        short_layer[:18] = b"Geomap v3.60 Layer"
        layer_name = "比例尺".encode("gb18030")
        struct.pack_into("<I", short_layer, 0x362, len(layer_name))
        short_layer[0x366:0x366 + len(layer_name)] = layer_name
        content.extend(short_layer)

        chunks = self.parser.parse("长612砂体厚度图.GDB", bytes(content))

        text = "\n".join(chunk.text for chunk in chunks)
        self.assertTrue(all(len(chunk.text) <= CHUNK_SIZE for chunk in chunks))
        self.assertIn("图层数: 1", text)
        self.assertIn("图层名称: 比例尺", text)
        self.assertIn("属性记录数: 0", text)

    def test_gdb_binary_file_remains_parsable_but_is_not_uploadable(self):
        chunks = self.parser.parse("长63渗透率.gdb", b"\x00\x01permeability=12.5\x00\x02")

        self.assertFalse(self.parser.supports("长63渗透率.GDB"))
        self.assertEqual(len(chunks), 1)
        self.assertIn("长63渗透率.gdb", chunks[0].text)
        self.assertIn("permeability=12.5", chunks[0].text)

    def test_geomap_album_information_is_parsed_as_hierarchy(self):
        content = bytearray(2048)
        magic = b"GeoMap Album Information\n"
        content[:len(magic)] = magic

        def append_u32(value):
            content.extend(struct.pack("<I", value))

        def append_string(value):
            encoded = value.encode("gb18030")
            append_u32(len(encoded))
            content.extend(encoded)

        append_string("黑山梁数据体")
        append_u32(0)
        append_string("黑山梁数据体")
        append_u32(1)
        append_u32(3)
        append_string("渗透率")
        append_u32(1)
        append_u32(1)
        append_string("长63渗透率")
        append_string("MAPS")
        append_string("长63渗透率.GDB")

        chunks = self.parser.parse("黑山梁数据体.LST", bytes(content))
        text = "\n".join(chunk.text for chunk in chunks)

        self.assertIn("GeoMap Album Information", text)
        self.assertIn("图册名: 黑山梁数据体", text)
        self.assertIn("分类: 渗透率", text)
        self.assertIn("完整相对路径: MAPS/长63渗透率.GDB", text)

    def test_xls_uses_excel_compatibility_fallback_when_xlrd_cannot_decode(self):
        fallback = [(None, "年度数据", "井号\t产量")]
        decode_error = UnicodeDecodeError("utf-16-le", b"\x00\xd8", 0, 2, "illegal encoding")

        with (
            patch("xlrd.open_workbook", side_effect=decode_error),
            patch.object(self.parser, "_parse_xls_with_excel", return_value=fallback) as excel_fallback,
        ):
            chunks = self.parser.parse("2008.12.30.xls", b"legacy workbook")

        excel_fallback.assert_called_once_with(b"legacy workbook")
        self.assertIsNone(chunks[0].section_path)
        self.assertIn("井号 产量", chunks[0].text)

    def test_uses_only_size_and_overlap_across_source_boundaries(self):
        chunks = self.parser._chunk_sources([
            (1, "第一页", "A" * 250),
            (2, "第二页", "B" * 147 + " " + "B" * 152),
        ])

        self.assertEqual([len(chunk.text) for chunk in chunks], [CHUNK_SIZE, 202])
        self.assertEqual(chunks[0].text[-CHUNK_OVERLAP:], chunks[1].text[:CHUNK_OVERLAP])
        self.assertTrue(all(chunk.page_number is None for chunk in chunks))
        self.assertTrue(all(chunk.section_path is None for chunk in chunks))

    def test_uploadable_formats_are_restricted(self):
        allowed = [
            "records.jsonl",
            "data.JSON",
            "report.pdf",
            "report.docx",
            "table.xlsx",
            "slides.pptx",
            "notes.txt",
            "readme.md",
            "readme.markdown",
        ]
        rejected = ["legacy.doc", "table.xls", "data.csv", "map.gdb", "README"]

        self.assertTrue(all(self.parser.supports(name) for name in allowed))
        self.assertTrue(all(not self.parser.supports(name) for name in rejected))

    def test_jsonl_is_parsed_as_individual_json_records(self):
        chunks = self.parser.parse(
            "records.jsonl",
            '{"井号":"化31"}\n{"产量":12}\n'.encode("utf-8"),
        )

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].text, '{"井号": "化31"}\n{"产量": 12}')


if __name__ == "__main__":
    unittest.main()
